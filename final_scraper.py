#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Claude Code中文教程完整爬虫
抓取 https://claudecode.tangshuang.net 的所有章节内容
"""

import asyncio
import json
import re
from pathlib import Path
from urllib.parse import urljoin

from playwright.async_api import async_playwright
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ClaudeCodeTutorialScraper:
    def __init__(self):
        self.base_url = "https://claudecode.tangshuang.net"
        self.tutorial_links = []
        self.scraped_content = []
        self.progress_file = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据/scraping_progress.json")

    def load_progress(self):
        """加载进度"""
        if self.progress_file.exists():
            with open(self.progress_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {'completed': [], 'failed': []}

    def save_progress(self, progress):
        """保存进度"""
        with open(self.progress_file, 'w', encoding='utf-8') as f:
            json.dump(progress, f, ensure_ascii=False, indent=2)

    async def get_all_tutorial_links(self, page):
        """获取所有教程链接"""
        print("正在获取所有教程链接...")

        # 访问第一个章节页面（这个页面包含完整的目录）
        await page.goto(f"{self.base_url}/course/1.1%20Claude%20Code%E6%98%AF%E4%BB%80%E4%B9%88", wait_until="networkidle")
        await page.wait_for_timeout(3000)

        # 提取所有课程链接
        links = await page.evaluate("""() => {
            const links = Array.from(document.querySelectorAll('a[href*="/course/"]'));
            return links.map(link => ({
                text: link.textContent.trim(),
                href: link.getAttribute('href')
            })).filter(link => link.text && link.href);
        }""")

        # 去重并排序
        seen = set()
        unique_links = []
        for link in links:
            if link['href'] not in seen:
                seen.add(link['href'])
                unique_links.append(link)

        # 按章节号排序
        def extract_chapter_num(href):
            match = re.search(r'/course/(\d+)\.', href)
            return int(match.group(1)) if match else 0

        unique_links.sort(key=lambda x: extract_chapter_num(x['href']))

        print(f"找到 {len(unique_links)} 个教程链接")
        return unique_links

    async def fetch_page_content(self, page, link_info, progress):
        """抓取单个页面内容"""
        url = urljoin(self.base_url, link_info['href'])

        # 跳过已完成的
        if url in progress['completed']:
            print(f"  ✓ 已缓存，跳过")
            return None

        try:
            await page.goto(url, wait_until="networkidle", timeout=30000)
            await page.wait_for_timeout(1500)

            # 提取内容
            content = await page.evaluate("""() => {
                // 查找主要内容区域
                const article = document.querySelector('article') ||
                               document.querySelector('main') ||
                               document.querySelector('[class*="content"]') ||
                               document.querySelector('[class*="markdown"]') ||
                               document.body;

                if (!article) return null;

                // 提取标题
                const title = document.querySelector('h1')?.textContent.trim() ||
                             document.querySelector('h2')?.textContent.trim() ||
                             '';

                // 提取所有内容
                const elements = article.querySelectorAll('h1, h2, h3, h4, h5, h6, p, pre, code, ul, ol, blockquote');

                const structuredContent = [];

                elements.forEach(el => {
                    const tagName = el.tagName.toLowerCase();

                    if (['h1', 'h2', 'h3', 'h4', 'h5', 'h6'].includes(tagName)) {
                        structuredContent.push({
                            type: 'heading',
                            level: parseInt(tagName[1]),
                            text: el.textContent.trim()
                        });
                    } else if (tagName === 'p') {
                        const text = el.textContent.trim();
                        if (text && text.length > 5) {
                            structuredContent.push({
                                type: 'paragraph',
                                text: text
                            });
                        }
                    } else if (tagName === 'pre' || tagName === 'code') {
                        const text = el.textContent.trim();
                        if (text) {
                            structuredContent.push({
                                type: 'code',
                                text: text
                            });
                        }
                    } else if (tagName === 'ul' || tagName === 'ol') {
                        const items = Array.from(el.querySelectorAll('li')).map(li => li.textContent.trim());
                        if (items.length > 0) {
                            structuredContent.push({
                                type: 'list',
                                ordered: tagName === 'ol',
                                items: items
                            });
                        }
                    } else if (tagName === 'blockquote') {
                        const text = el.textContent.trim();
                        if (text) {
                            structuredContent.push({
                                type: 'quote',
                                text: text
                            });
                        }
                    }
                });

                return {
                    title: title,
                    content: structuredContent
                };
            }""")

            if content and content['content']:
                result = {
                    'title': link_info['text'],
                    'page_title': content['title'],
                    'url': url,
                    'content': content['content']
                }

                progress['completed'].append(url)
                self.save_progress(progress)

                return result
            else:
                progress['failed'].append(url)
                self.save_progress(progress)
                return None

        except Exception as e:
            print(f"    错误: {str(e)}")
            progress['failed'].append(url)
            self.save_progress(progress)
            return None

    async def scrape_all_tutorials(self):
        """抓取所有教程"""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()

            # 获取所有链接
            self.tutorial_links = await self.get_all_tutorial_links(page)

            # 加载进度
            progress = self.load_progress()

            print(f"\n开始抓取 {len(self.tutorial_links)} 个页面...")
            print(f"已缓存: {len(progress['completed'])} 个")
            print(f"待抓取: {len(self.tutorial_links) - len(progress['completed'])} 个\n")

            # 抓取每个页面
            success_count = 0
            for idx, link in enumerate(self.tutorial_links, 1):
                print(f"[{idx}/{len(self.tutorial_links)}] {link['text']}")
                print(f"  URL: {link['href']}")

                content = await self.fetch_page_content(page, link, progress)

                if content:
                    self.scraped_content.append(content)
                    success_count += 1
                    print(f"  ✓ 成功 ({len(content['content'])} 个内容块)")
                else:
                    print(f"  ✗ 失败或为空")

                # 每抓取10个页面保存一次
                if idx % 10 == 0:
                    print(f"\n--- 进度保存 ({idx}/{len(self.tutorial_links)}) ---\n")

            await browser.close()

            print(f"\n抓取完成！")
            print(f"成功: {success_count}/{len(self.tutorial_links)}")
            print(f"失败: {len(progress['failed'])}")
            print(f"缓存: {len(progress['completed']) - success_count}")

            return self.scraped_content

    def organize_by_chapter(self):
        """按章节组织内容"""
        chapters = {}

        for item in self.scraped_content:
            # 提取章节号
            match = re.search(r'(\d+)\.(.+)', item['title'])
            if match:
                chapter_num = match.group(1)
                section_title = match.group(2).strip()

                if chapter_num not in chapters:
                    chapters[chapter_num] = {
                        'chapter_num': chapter_num,
                        'sections': []
                    }

                chapters[chapter_num]['sections'].append({
                    'title': section_title,
                    'page_title': item['page_title'],
                    'url': item['url'],
                    'content': item['content']
                })

        # 按章节号排序
        sorted_chapters = [chapters[k] for k in sorted(chapters.keys(), key=lambda x: int(x))]
        return sorted_chapters

    def save_to_markdown(self, output_path):
        """保存为Markdown文件"""
        print(f"\n正在保存Markdown文件: {output_path}")

        chapters = self.organize_by_chapter()

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Claude Code 中文教程\n\n")
            f.write("本教程从 https://claudecode.tangshuang.net 抓取\n\n")
            f.write(f"共 {len(chapters)} 章，{len(self.scraped_content)} 个小节\n\n")
            f.write("---\n\n")

            for chapter in chapters:
                chapter_num = chapter['chapter_num']
                sections = chapter['sections']

                f.write(f"# 第 {chapter_num} 章\n\n")

                for section in sections:
                    f.write(f"## {section['title']}\n\n")
                    f.write(f"**URL**: {section['url']}\n\n")

                    # 写入内容
                    for item in section['content']:
                        if item['type'] == 'heading':
                            level = min(item['level'] + 2, 6)
                            f.write(f"{'#' * level} {item['text']}\n\n")
                        elif item['type'] == 'paragraph':
                            f.write(f"{item['text']}\n\n")
                        elif item['type'] == 'code':
                            f.write(f"```\n{item['text']}\n```\n\n")
                        elif item['type'] == 'list':
                            for list_item in item['items']:
                                prefix = f"{item['ordered']}." if item.get('ordered') else "-"
                                f.write(f"{prefix} {list_item}\n")
                            f.write("\n")
                        elif item['type'] == 'quote':
                            f.write(f"> {item['text']}\n\n")

                    f.write("---\n\n")

        print(f"✓ Markdown文件已保存")

    def save_to_word(self, output_path):
        """保存为Word文件"""
        print(f"\n正在保存Word文件: {output_path}")

        chapters = self.organize_by_chapter()
        doc = Document()

        # 主标题
        title = doc.add_heading('Claude Code 中文教程', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 说明
        doc.add_paragraph('本教程从 https://claudecode.tangshuang.net 抓取')
        doc.add_paragraph(f'共 {len(chapters)} 章，{len(self.scraped_content)} 个小节')
        doc.add_paragraph()

        for chapter in chapters:
            chapter_num = chapter['chapter_num']
            sections = chapter['sections']

            # 章节标题
            doc.add_heading(f'第 {chapter_num} 章', level=1)

            for section in sections:
                # 小节标题
                doc.add_heading(section['title'], level=2)

                # 内容
                for item in section['content']:
                    if item['type'] == 'heading':
                        level = min(item['level'] + 1, 3)
                        doc.add_heading(item['text'], level=level)
                    elif item['type'] == 'paragraph':
                        doc.add_paragraph(item['text'])
                    elif item['type'] == 'code':
                        # 代码块使用等宽字体
                        para = doc.add_paragraph(item['text'])
                        para.style = 'Code'
                    elif item['type'] == 'list':
                        for list_item in item['items']:
                            doc.add_paragraph(list_item, style='List Number' if item.get('ordered') else 'List Bullet')
                    elif item['type'] == 'quote':
                        para = doc.add_paragraph(item['text'])
                        para.style = 'Quote'

                # 分页（除了最后一章的最后一节）
                doc.add_page_break()

        doc.save(output_path)
        print(f"✓ Word文件已保存")

    def save_summary(self, output_path):
        """保存抓取摘要"""
        chapters = self.organize_by_chapter()

        summary = {
            'total_chapters': len(chapters),
            'total_sections': len(self.scraped_content),
            'chapters': []
        }

        for chapter in chapters:
            summary['chapters'].append({
                'chapter_num': chapter['chapter_num'],
                'sections_count': len(chapter['sections']),
                'sections': [s['title'] for s in chapter['sections']]
            })

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)

        print(f"✓ 摘要已保存: {output_path}")


async def main():
    """主函数"""
    scraper = ClaudeCodeTutorialScraper()

    print("=" * 70)
    print("Claude Code中文教程完整爬虫")
    print("=" * 70)
    print()

    # 抓取所有内容
    await scraper.scrape_all_tutorials()

    if not scraper.scraped_content:
        print("没有抓取到任何内容，程序退出")
        return

    # 保存文件
    base_path = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据")
    scraper.save_to_markdown(base_path / "Claude_Code中文教程_完整版.md")
    scraper.save_to_word(base_path / "Claude_Code中文教程_完整版.docx")
    scraper.save_summary(base_path / "tutorial_summary.json")

    print("\n" + "=" * 70)
    print("所有任务完成！")
    print("=" * 70)


if __name__ == "__main__":
    asyncio.run(main())

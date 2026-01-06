#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从已抓取的数据生成Word文档
"""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def load_scraped_data():
    """加载已抓取的数据"""
    progress_file = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据/scraping_progress.json")
    links_file = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据/course_links.json")

    with open(progress_file, 'r', encoding='utf-8') as f:
        progress = json.load(f)

    with open(links_file, 'r', encoding='utf-8') as f:
        all_links = json.load(f)

    return progress, all_links


def organize_by_chapter(progress, all_links):
    """按章节组织内容"""
    chapters = {}

    for url in progress['completed']:
        # 找到对应的链接信息
        link_info = None
        for link in all_links:
            full_url = f"https://claudecode.tangshuang.net{link['href']}"
            if full_url == url and link['text']:
                link_info = link
                break

        if not link_info:
            continue

        # 提取章节号
        match = re.search(r'(\d+)\.(.+)', link_info['text'])
        if match:
            chapter_num = match.group(1)
            section_title = match.group(2).strip()

            if chapter_num not in chapters:
                chapters[chapter_num] = {
                    'chapter_num': chapter_num,
                    'sections': []
                }

            chapters[chapter_num]['sections'].append({
                'title': section_title,
                'url': url
            })

    # 按章节号排序
    sorted_chapters = [chapters[k] for k in sorted(chapters.keys(), key=lambda x: int(x))]
    return sorted_chapters


def generate_word_from_markdown():
    """从Markdown文件生成Word文档"""
    print("正在读取Markdown文件...")

    md_file = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据/Claude_Code中文教程_完整版.md")
    output_file = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据/Claude_Code中文教程_完整版.docx")

    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print("正在创建Word文档...")
    doc = Document()

    # 主标题
    title = doc.add_heading('Claude Code 中文教程', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 处理Markdown内容
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()

            if level <= 3:
                doc.add_heading(title_text, level=level)
            else:
                doc.add_paragraph(title_text)
        # 处理代码块
        elif line.startswith('```'):
            # 跳过代码块标记
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            code_text = '\n'.join(code_lines)
            if code_text.strip():
                para = doc.add_paragraph(code_text)
                # 设置等宽字体
                if para.runs:
                    run = para.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        # 处理引用
        elif line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            para = doc.add_paragraph(quote_text)
            para.style = 'Quote'
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line.lstrip('-*').strip()
            doc.add_paragraph(list_text, style='List Bullet')
        elif re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.', '', line).strip()
            doc.add_paragraph(list_text, style='List Number')
        # 处理粗体
        elif '**' in line:
            # 简单处理，添加为普通段落
            clean_text = line.replace('**', '').strip()
            doc.add_paragraph(clean_text)
        # 处理URL行
        elif line.startswith('**URL**:') or line.startswith('*URL*:'):
            # 跳过这些元数据行
            pass
        # 处理分隔线
        elif line.startswith('---'):
            pass
        # 普通段落
        else:
            if line.strip():
                doc.add_paragraph(line)

        i += 1

    print("正在保存Word文档...")
    doc.save(output_file)
    print(f"✓ Word文档已保存: {output_file}")
    print(f"  文件大小: {output_file.stat().st_size / 1024 / 1024:.2f} MB")


def main():
    print("=" * 70)
    print("从Markdown生成Word文档")
    print("=" * 70)
    print()

    generate_word_from_markdown()

    print("\n" + "=" * 70)
    print("完成！")
    print("=" * 70)


if __name__ == "__main__":
    main()

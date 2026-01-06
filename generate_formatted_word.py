#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从markdown生成排版优化的Word文档
使用python-docx库
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_word(md_file, docx_file):
    """将markdown文件转换为格式化的Word文档"""

    print(f"读取markdown文件: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    print(f"总行数: {len(lines)}")

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    print("解析markdown并构建Word文档...")

    # 状态变量
    in_code_block = False
    code_lines = []
    list_level = 0

    for i, line in enumerate(lines):
        line = line.rstrip('\n')

        # 代码块处理
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.style = 'No Spacing'
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # 空行
        if not line.strip():
            if i < len(lines) - 1 and lines[i + 1].strip():
                doc.add_paragraph()  # 只在需要时添加空行
            continue

        # 一级标题（章）
        if line.startswith('# 第 ') and ' 章' in line:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(line.strip('# ').strip())
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.space_after = Pt(12)
            continue

        # 二级标题（节）
        if line.startswith('## ') and not line.startswith('###'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('# ').strip())
            run.font.size = Pt(14)
            run.font.bold = True
            p.space_before = Pt(12)
            p.space_after = Pt(6)
            continue

        # 三级标题（小节）
        if line.startswith('### ') and not line.startswith('####'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('# ').strip())
            run.font.size = Pt(12)
            run.font.bold = True
            p.space_before = Pt(6)
            p.space_after = Pt(3)
            continue

        # 四级标题
        if line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('# ').strip())
            run.font.size = Pt(11)
            run.font.bold = True
            p.space_before = Pt(3)
            p.space_after = Pt(3)
            continue

        # 列表项
        if line.strip().startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.', line.strip()):
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            p.style.font.size = Pt(11)
            continue

        # URL引用
        if line.startswith('**URL**:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.italic = True
            continue

        # 普通段落
        # 处理行内代码
        if '`' in line:
            p = doc.add_paragraph()
            parts = line.split('`')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # 普通文本
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                else:
                    # 代码
                    run = p.add_run(part)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.style.font.size = Pt(11)

    print("保存Word文档...")
    doc.save(docx_file)
    print(f"✓ 已保存: {docx_file}")

    return True

def main():
    md_file = "/Users/lartemacfiles/Desktop/VPT-初诊数据/Claude_Code中文教程_完整版.md"
    docx_file = "/Users/lartemacfiles/Desktop/VPT-初诊数据/Claude_Code中文教程_完整版_优化版.docx"

    print("=" * 80)
    print("生成排版优化的Word文档")
    print("=" * 80)

    try:
        success = parse_markdown_to_word(md_file, docx_file)

        if success:
            print("\n" + "=" * 80)
            print("完成！")
            print("=" * 80)
            print(f"\n输入文件: {md_file}")
            print(f"输出文件: {docx_file}")
    except Exception as e:
        print(f"✗ 生成失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Claude Code中文教程爬虫 V2
抓取 https://claudecode.tangshuang.net 的所有章节内容
"""

import asyncio
import json
import re
from pathlib import Path
from urllib.parse import urljoin, unquote

from playwright.async_api import async_playwright
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ClaudeCodeScraperV2:
    def __init__(self):
        self.base_url = "https://claudecode.tangshuang.net"
        self.chapters = []
        self.full_content = []

    async def extract_all_links(self, page):
        """从页面中提取所有教程链接"""
        print("正在提取所有教程链接...")

        # 等待页面完全加载
        await page.wait_for_timeout(5000)

        # 尝试点击"开始免费阅读"或类似按钮
        try:
            # 尝试查找并点击阅读按钮
            buttons = await page.locator('a, button').all()
            for button in buttons:
                text = await button.text_content()
                if text and any(keyword in text for keyword in ['阅读', '教程', '开始', 'Read']):
                    await button.click()
                    await page.wait_for_timeout(2000)
                    break
        except:
            pass

        # 获取所有链接
        content = await page.content()
        soup = BeautifulSoup(content, 'html.parser')

        tutorial_links = []

        # 查找所有a标签
        for link in soup.find_all('a', href=True):
            href = link['href']
            text = link.get_text(strip=True)

            # 过滤出教程相关链接
            if href.startswith('/course/') or 'course' in href:
                full_url = urljoin(self.base_url, href)

                # 清理URL
                if full_url not in [x['url'] for x in tutorial_links]:
                    tutorial_links.append({
                        'title': text,
                        'url': full_url
                    })

        print(f"  找到 {len(tutorial_links)} 个教程链接")
        return tutorial_links

    async def discover_structure(self, page):
        """发现完整的教程结构"""
        print("正在发现教程结构...")

        # 访问首页
        await page.goto(self.base_url, wait_until="networkidle", timeout=30000)

        # 提取链接
        links = await self.extract_all_links(page)

        # 如果没有找到足够的链接，尝试直接访问常见的章节URL
        if len(links) < 30:
            print("  链接数量不足，尝试访问章节URL模式...")
            links = []

            # 尝试不同的URL模式
            patterns = [
                '/course/{}',
                '/chapter-{}',
                '/{}',
                '/course/chapter-{}',
            ]

            for pattern in patterns:
                for i in range(1, 35):
                    url = f"{self.base_url}{pattern.format(i)}"
                    links.append({
                        'title': f'第{i}章',
                        'url': url
                    })

        return links

    async def fetch_chapter_content(self, page, url, title=""):
        """获取章节内容"""
        try:
            await page.goto(url, wait_until="networkidle", timeout=30000)
            await page.wait_for_timeout(2000)

            # 获取页面HTML
            content = await page.content()
            soup = BeautifulSoup(content, 'html.parser')

            # 提取标题
            h1 = soup.find('h1')
            h2 = soup.find('h2')
            page_title = h1.get_text(strip=True) if h1 else (h2.get_text(strip=True) if h2 else title)

            # 尝试多种方式提取主要内容
            main_content = None

            # 方法1: 查找article标签
            article = soup.find('article')
            if article:
                main_content = article

            # 方法2: 查找main标签
            if not main_content:
                main = soup.find('main')
                if main:
                    main_content = main

            # 方法3: 查找特定的div
            if not main_content:
                # 尝试找到包含最多文本的div
                divs = soup.find_all('div')
                if divs:
                    main_content = max(divs, key=lambda d: len(d.get_text()))

            if main_content:
                # 提取结构化内容
                structured_content = self.extract_structured_content(main_content)

                return {
                    'url': url,
                    'title': page_title,
                    'content': structured_content,
                    'raw_html': str(main_content)
                }
            else:
                return None

        except Exception as e:
            print(f"    错误: {str(e)}")
            return None

    def extract_structured_content(self, element):
        """提取结构化内容，保留标题层级"""
        content = []
        current_element = element

        # 移除脚本和样式标签
        for script in current_element(['script', 'style', 'nav', 'header', 'footer']):
            script.decompose()

        # 遍历所有子元素
        for child in current_element.descendants:
            if child.name:
                # 处理标题
                if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(child.name[1])
                    text = child.get_text(strip=True)
                    if text:
                        content.append({
                            'type': 'heading',
                            'level': level,
                            'text': text
                        })

                # 处理段落
                elif child.name == 'p':
                    text = child.get_text(strip=True)
                    if text and len(text) > 5:  # 过滤太短的内容
                        content.append({
                            'type': 'paragraph',
                            'text': text
                        })

                # 处理代码块
                elif child.name in ['pre', 'code']:
                    text = child.get_text(strip=True)
                    if text:
                        content.append({
                            'type': 'code',
                            'text': text
                        })

                # 处理列表
                elif child.name in ['ul', 'ol']:
                    items = child.find_all('li', recursive=False)
                    list_items = []
                    for item in items:
                        text = item.get_text(strip=True)
                        if text:
                            list_items.append(text)

                    if list_items:
                        content.append({
                            'type': 'list',
                            'items': list_items
                        })

        return content

    async def scrape_all(self):
        """抓取所有内容"""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()

            # 发现结构
            all_links = await self.discover_structure(page)
            print(f"\n总共需要抓取 {len(all_links)} 个页面")

            # 抓取每个章节
            success_count = 0
            for idx, link in enumerate(all_links, 1):
                print(f"\n进度: {idx}/{len(all_links)} - {link['title']}")
                print(f"  URL: {link['url']}")

                content = await self.fetch_chapter_content(page, link['url'], link['title'])

                if content and content['content']:
                    content['chapter_title'] = link['title']
                    content['chapter_num'] = idx
                    self.full_content.append(content)
                    success_count += 1
                    print(f"  ✓ 成功抓取 ({len(content['content'])} 个内容块)")
                else:
                    print(f"  ✗ 抓取失败")

            await browser.close()

        print(f"\n成功抓取 {success_count}/{len(all_links)} 个章节")
        return self.full_content

    def save_to_markdown(self, output_path):
        """保存为Markdown文件"""
        print(f"\n正在保存Markdown文件: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Claude Code 中文教程\n\n")
            f.write("本教程从 https://claudecode.tangshuang.net 抓取\n\n")
            f.write("---\n\n")

            for content in self.full_content:
                f.write(f"## {content['chapter_title']}\n\n")
                f.write(f"### {content['title']}\n\n")

                # 写入结构化内容
                for item in content['content']:
                    if item['type'] == 'heading':
                        level = item['level'] + 2  # 调整层级
                        f.write(f"{'#' * level} {item['text']}\n\n")
                    elif item['type'] == 'paragraph':
                        f.write(f"{item['text']}\n\n")
                    elif item['type'] == 'code':
                        f.write(f"```\n{item['text']}\n```\n\n")
                    elif item['type'] == 'list':
                        for list_item in item['items']:
                            f.write(f"- {list_item}\n")
                        f.write("\n")

                f.write("---\n\n")

        print("✓ Markdown文件已保存")

    def save_to_word(self, output_path):
        """保存为Word文件"""
        print(f"\n正在保存Word文件: {output_path}")
        doc = Document()

        # 主标题
        title = doc.add_heading('Claude Code 中文教程', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 说明
        doc.add_paragraph('本教程从 https://claudecode.tangshuang.net 抓取')
        doc.add_paragraph()

        for content in self.full_content:
            # 章节标题
            doc.add_heading(content['chapter_title'], 1)

            # 页面标题
            doc.add_heading(content['title'], 2)

            # 内容
            for item in content['content']:
                if item['type'] == 'heading':
                    doc.add_heading(item['text'], level=min(item['level'] + 1, 3))
                elif item['type'] == 'paragraph':
                    doc.add_paragraph(item['text'])
                elif item['type'] == 'code':
                    doc.add_paragraph(item['text'], style='Code')
                elif item['type'] == 'list':
                    for list_item in item['items']:
                        doc.add_paragraph(list_item, style='List Bullet')

            doc.add_page_break()

        doc.save(output_path)
        print("✓ Word文件已保存")

    def save_summary(self, output_path):
        """保存抓取摘要"""
        summary = {
            'total_chapters': len(self.full_content),
            'chapters': []
        }

        for content in self.full_content:
            summary['chapters'].append({
                'chapter_num': content['chapter_num'],
                'chapter_title': content['chapter_title'],
                'title': content['title'],
                'url': content['url'],
                'content_blocks': len(content['content'])
            })

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)

        print(f"✓ 摘要已保存: {output_path}")


async def main():
    """主函数"""
    scraper = ClaudeCodeScraperV2()

    print("=" * 60)
    print("Claude Code中文教程爬虫 V2")
    print("=" * 60)

    # 抓取所有内容
    await scraper.scrape_all()

    # 保存文件
    base_path = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据")
    scraper.save_to_markdown(base_path / "Claude_Code中文教程_v2.md")
    scraper.save_to_word(base_path / "Claude_Code中文教程_v2.docx")
    scraper.save_summary(base_path / "scrape_summary_v2.json")

    print("\n" + "=" * 60)
    print(f"抓取完成! 共获取 {len(scraper.full_content)} 个章节")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())

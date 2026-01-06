#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Claude Code中文教程爬虫
抓取 https://claudecode.tangshuang.net 的所有章节内容
"""

import asyncio
import json
import re
from pathlib import Path
from urllib.parse import urljoin

from playwright.async_api import async_playwright
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ClaudeCodeScraper:
    def __init__(self):
        self.base_url = "https://claudecode.tangshuang.net"
        self.chapters = []
        self.full_content = []

    async def fetch_page(self, page, url):
        """获取单个页面的内容"""
        print(f"正在访问: {url}")
        try:
            await page.goto(url, wait_until="networkidle", timeout=30000)
            await page.wait_for_timeout(2000)  # 等待JavaScript渲染

            # 获取页面内容
            content = await page.content()
            soup = BeautifulSoup(content, 'html.parser')

            # 尝试多种方式提取主要内容
            main_content = None

            # 方法1: 查找article标签
            article = soup.find('article')
            if article:
                main_content = article
            else:
                # 方法2: 查找main标签
                main = soup.find('main')
                if main:
                    main_content = main
                else:
                    # 方法3: 查找包含内容的div
                    divs = soup.find_all('div', class_=re.compile(r'content|markdown|article'))
                    if divs:
                        main_content = divs[0]

            if main_content:
                # 提取标题
                title = soup.find('h1')
                title_text = title.get_text(strip=True) if title else "未找到标题"

                # 提取所有文本内容
                text_content = main_content.get_text(separator='\n', strip=True)

                return {
                    'url': url,
                    'title': title_text,
                    'content': text_content,
                    'html': str(main_content)
                }
            else:
                print(f"  警告: 未找到主要内容")
                return None

        except Exception as e:
            print(f"  错误: {str(e)}")
            return None

    async def discover_chapters(self, page):
        """发现所有章节链接"""
        print("正在发现章节链接...")
        await page.goto(self.base_url, wait_until="networkidle", timeout=30000)
        await page.wait_for_timeout(3000)

        # 尝试从页面中提取所有链接
        content = await page.content()
        soup = BeautifulSoup(content, 'html.parser')

        # 查找所有可能包含章节导航的元素
        nav_links = []

        # 查找所有a标签
        for link in soup.find_all('a', href=True):
            href = link['href']
            text = link.get_text(strip=True)

            # 过滤出章节链接
            if any(keyword in text for keyword in ['第', '章', 'Chapter']):
                full_url = urljoin(self.base_url, href)
                nav_links.append({
                    'title': text,
                    'url': full_url
                })

        # 如果没有找到，尝试常见的URL模式
        if not nav_links:
            print("  未从页面中发现章节链接，尝试URL模式...")
            for i in range(1, 35):  # 34章
                nav_links.append({
                    'title': f'第{i}章',
                    'url': f'{self.base_url}/{i}'
                })

        return nav_links

    async def scrape_all(self):
        """抓取所有章节"""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()

            # 发现章节
            chapters = await self.discover_chapters(page)
            print(f"发现 {len(chapters)} 个章节")

            # 抓取每个章节
            for idx, chapter in enumerate(chapters, 1):
                print(f"\n进度: {idx}/{len(chapters)}")
                content = await self.fetch_page(page, chapter['url'])

                if content:
                    content['chapter_title'] = chapter['title']
                    content['chapter_num'] = idx
                    self.full_content.append(content)

            await browser.close()

        return self.full_content

    def save_to_markdown(self, output_path):
        """保存为Markdown文件"""
        print(f"\n正在保存Markdown文件: {output_path}")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Claude Code 中文教程\n\n")
            f.write("本教程从 https://claudecode.tangshuang.net 抓取\n\n")
            f.write("---\n\n")

            for content in self.full_content:
                f.write(f"## {content['chapter_title']}\n\n")
                if content['title'] != content['chapter_title']:
                    f.write(f"### {content['title']}\n\n")
                f.write(f"{content['content']}\n\n")
                f.write("---\n\n")

        print(f"Markdown文件已保存")

    def save_to_word(self, output_path):
        """保存为Word文件"""
        print(f"\n正在保存Word文件: {output_path}")
        doc = Document()

        # 标题
        title = doc.add_heading('Claude Code 中文教程', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 说明
        doc.add_paragraph('本教程从 https://claudecode.tangshuang.net 抓取')
        doc.add_paragraph()

        for content in self.full_content:
            # 章节标题
            doc.add_heading(content['chapter_title'], 1)

            # 小标题
            if content['title'] != content['chapter_title']:
                doc.add_heading(content['title'], 2)

            # 内容段落
            paragraphs = content['content'].split('\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

            doc.add_page_break()

        doc.save(output_path)
        print(f"Word文件已保存")

    def save_summary(self, output_path):
        """保存抓取摘要"""
        summary = {
            'total_chapters': len(self.full_content),
            'chapters': []
        }

        for content in self.full_content:
            summary['chapters'].append({
                'chapter_num': content['chapter_num'],
                'chapter_title': content['chapter_title'],
                'title': content['title'],
                'url': content['url'],
                'content_length': len(content['content'])
            })

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)

        print(f"摘要已保存: {output_path}")


async def main():
    """主函数"""
    scraper = ClaudeCodeScraper()

    print("=" * 60)
    print("Claude Code中文教程爬虫")
    print("=" * 60)

    # 抓取所有内容
    await scraper.scrape_all()

    # 保存文件
    base_path = Path("/Users/lartemacfiles/Desktop/VPT-初诊数据")
    scraper.save_to_markdown(base_path / "Claude_Code中文教程.md")
    scraper.save_to_word(base_path / "Claude_Code中文教程.docx")
    scraper.save_summary(base_path / "scrape_summary.json")

    print("\n" + "=" * 60)
    print(f"抓取完成! 共获取 {len(scraper.full_content)} 个章节")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
